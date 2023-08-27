from flask import Flask, request
from docx import Document
import json

app = Flask(__name__)

@app.route('/createdoc', methods=['POST'])
def create_file():
 try:
    # Get the JSON data from the request
    json_data = request.get_json()
    # Convert JSON data to dictionary
    variables = json.loads(json.dumps(json_data))
    
    # Specify input and output file paths
    template_file_path = 'Proposal_Letter.docx'
    output_file_path = 'result.docx'
    
    # Function to replace text in a paragraph
    def replace_text_in_paragraph(paragraph, key, value):
      if key in paragraph.text:
       paragraph.text = paragraph.text.replace(key, value)
    
    # Load the template document
    template_document = Document(template_file_path)
    
    # Replace variables in paragraphs
    for variable_key, variable_value in variables.items():
      for paragraph in template_document.paragraphs:
       replace_text_in_paragraph(paragraph, variable_key, variable_value)
    
    # Replace variables in table cells
    for variable_key, variable_value in variables.items():
      table = template_document.tables[0]
      for row in table.rows:
        for cell in row.cells:
          if variable_key in cell.text:
                cell.text = variable_value
    
    # Save the updated document
    template_document.save(output_file_path)
    
    # Return the output file as a download
    return ('')

 except:
# Return an error message if something goes wrong
   return 'Error creating document', 500

if __name__ =="__main__":
  app.run(host='0.0.0.0',debug=True)
